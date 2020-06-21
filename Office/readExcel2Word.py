# -*- coding: utf-8 -*-
"""

first：
pip install openpyxl
pip install python_docx

读取原文档: A-‘英文名’; B-‘简写’; C-‘中文名’; D-‘注释’。

拆分出‘英文名’、‘简写’、‘中文名’、‘注释’然后组合成所需样式：
‘英文名’(‘简写’)  ‘中文名’  ‘注释’
‘简写’  ‘中文名’  ‘英文名’的缩写
最后统一进行排序

@author: guodxu@qq.com
"""


import os
import docx
import pandas as pd
import openpyxl as xl


dir=os.path.split(os.path.realpath(__file__))[0]
print(dir)
wbook=xl.load_workbook(os.path.join(dir, 'dict-full.xlsx'))
wsheet=wbook.active
print(wsheet.title)
print(wsheet.max_row, wsheet.max_column)


i=0
# columns[shortname, shortnameUpper, englisname, englisnameUpper, chinesename, description]
dtxt={}
df = pd.DataFrame(columns=['sn','su','en','eu','cn','txt'])

names={'auger', 'anderson', 'brillouin', 'bridgman', 'brinell', 'boltzmann', 'czochralski', 
       'faraday', 'fourier',  'frankel', 'gruneisen', 'josephson', 'joyce-dixon', 'kelvin',
       'kirchhoff', 'kirk', 'kondo', 'madelung', 'manchester', 'michelson'}
firstCh = {'是', '在', '对', '指', '其', '又', '根', '它', '这', '从', '该', '包', '每', '也', '通'}
for row in wsheet.rows:
    if not row[2].value:
        continue
    
    cnName=row[2].value.strip()
    if cnName[-3:] == "缩写。":  #"缩略语":
        continue
    
    enName=row[0].value.strip()
    shortName=row[1].value
    descrip=row[3].value    
        
    if(shortName):
        shortName=shortName.strip()
        shortName=shortName.strip('()')
    else:
        shortName=''
    
    descrip=descrip.strip()
    if(descrip[:len(cnName)] == cnName):
        descrip=descrip[len(cnName):]
    elif(descrip[:len(shortName)] == shortName):
        descrip=descrip[len(shortName):]
    if(len(descrip)<2):
        print(enName, cnName)
        continue
    
#    if(descrip[0] not in firstCh):
#        descrip='是' + descrip
    
    dtxt['sn']=shortName
    dtxt['su']=shortName.lower()
    
#    lst = []
#    enlist = enName.split()
#    for e in enlist:
#        en=e
#        if(e.istitle()):
#            en=str.lower(e)        
#        if(en in names):
#            en =str.title(e)
#        lst.append(en)
#    dtxt['en'] = ' '.join(lst)
    dtxt['en'] = enName
        
    dtxt['eu']=dtxt['en'].lower()
    dtxt['cn']=cnName
    dtxt['txt']=descrip
    df=df.append(dtxt, ignore_index=True)
#    print(dtxt)
#    i=i+1
#    if(i>10): 
#        break
    
    
    
#    
#    print(len(df))
#    print(enName, shortName, cnName, descrip)
#    i +=1
#    if(i>10):
#        break
#    
    
print("Total full-entry:", len(df))


# find the dumplicated
vcount = df.eu.value_counts()
print(vcount[vcount>1])


#############################################################################
# to excel
# columns=[sortname, title, description]
newdf = pd.DataFrame(columns=['sn','full', 'short', 'title','txt'])
dnew={}
for i in range(len(df)):
    # get full name
    rw = df.iloc[i]
    dnew['sn'] = rw.eu
    if(not str.isalpha(rw.eu[0])):
        dnew['sn'] = "zz" + rw.eu
    dnew['full']=rw.en
    short=''
    if(len(rw.sn)>0):
        short='('+rw.sn+')'
    dnew['short']=short
    dnew['title']=rw.cn
    
#    if(len(rw.sn)>0):
#        dnew['title'] = "{en} ({sn})  {cn}  ".format(en=rw.en, sn=rw.sn, cn=rw.cn)
#    else:
#        dnew['title'] = "{en}  {cn}  ".format(en=rw.en, cn=rw.cn)
#    
    dnew['txt'] = rw.txt
    newdf = newdf.append(dnew, ignore_index=True)
    
    # get short name
    su = rw.su.replace('.', '')
    if(len(su)>=3 and len(su)<6):
        dnew['sn'] = su
        if(not str.isalpha(rw.sn[0])):
            dnew['sn'] = "zz" + su   # 排在字母后面
        dnew['full']=rw.sn
        dnew['short']=''
        dnew['txt'] = ''
        dnew['title'] = "{en}的缩写。".format(en=rw.en)
        newdf = newdf.append(dnew, ignore_index=True)
        
print("Total: ", len(newdf), ", Full: ", len(df), ", Short: ", len(newdf)-len(df))


newdf.sort_values(by='sn', inplace=True)


# save to excel
nwb=xl.Workbook()
nsheet=nwb.create_sheet('sheet1', index=0)

for i in range(len(newdf)):
#for i in range(10):
    rw=newdf.iloc[i]
#    print(rw)
    nsheet.cell(row=i+1, column=1).value=rw.full
    nsheet.cell(row=i+1, column=2).value=rw.short
    nsheet.cell(row=i+1, column=3).value=rw.title
    nsheet.cell(row=i+1, column=4).value=rw.txt
    
nwb.save(os.path.join(dir,"dictFull-new.xlsx"))
print('save excel finish')


############################################################
# to word
# columns=[sortname, name, title, txt]
newdf = pd.DataFrame(columns=['sn','title','txt'])
dnew={}
for i in range(len(df)):
    # get full name
    rw = df.iloc[i]
    dnew['sn'] = rw.eu
    if(not str.isalpha(rw.eu[0])):
        dnew['sn'] = "zz" + rw.eu
    if(len(rw.sn)>0):
        dnew['title'] = "{en} ({sn})  {cn}  ".format(en=rw.en, sn=rw.sn, cn=rw.cn)
    else:
        dnew['title'] = "{en}  {cn}  ".format(en=rw.en, cn=rw.cn)
    dnew['txt'] = rw.txt
    newdf = newdf.append(dnew, ignore_index=True)
    
    # get short name
    su = rw.su.replace('.', '')
    if(len(su)>=3 and len(su)<6):
        dnew['sn'] = su
        if(not str.isalpha(rw.sn[0])):
            dnew['sn'] = "zz" + su   # 排在字母后面
        dnew['title'] = "{sn}  {cn}  ".format(sn=rw.sn, cn=rw.cn)
        dnew['txt'] = "{en}的缩写。".format(en=rw.en)
        newdf = newdf.append(dnew, ignore_index=True)

newdf.sort_values(by='sn', inplace=True)
        
# save to docx
hd = ''
doc=docx.Document()
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = docx.shared.Pt(12)
sty._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')

for i in range(len(newdf)):
    rw=newdf.iloc[i]
    if (hd != rw.sn[0]):
        if(len(hd)>0):
            doc.add_page_break()
        hd = rw.sn[0]
        doc.add_heading(hd.upper())
    p=doc.add_paragraph(style=sty)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = docx.shared.Pt(4)
    p.paragraph_format.space_after = 0
    
    run=p.add_run(rw.title)
    run.bold = True
    p.add_run(rw.txt)
    
    
doc.save(os.path.join(dir,"dictFull-new.docx"))
print('save word finish')