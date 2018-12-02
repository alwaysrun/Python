# -*- coding: utf-8 -*-
"""

first：
pip install python_docx

读取原文档: ACLS (Automatic Cassette Load Station )  真空机械手  是主要放置Cassette的地方。
拆分出‘英文名’、‘简写’、‘中文名’、‘注释’然后组合成所需样式：
‘英文名’(‘简写’)  ‘中文名’  ‘注释’
‘简写’  ‘中文名’  ‘英文名’的缩写
最后统一进行排序

@author: xugd(guodxu@qq.com)
"""

import os
import docx
import pandas as pd

dir=os.path.split(os.path.realpath(__file__))[0]
print(dir)
doc=docx.Document(os.path.join(dir,"dict-ori.docx"))

print("paragraphs: ", len(doc.paragraphs))

dtxt={}

# columns[shortname, shortnameUpper, englisname, englisnameUpper, chinesename, description]
df = pd.DataFrame(columns=['sn','su','en','eu','cn','txt'])
for par in doc.paragraphs:
    txt = par.text
    if(len(txt)<5):
        continue
    
    # get short name
    fi=txt.find('(')
    if(fi != -1):
        dtxt['sn'] = txt[:fi].strip()
        dtxt['su'] = str.upper(dtxt['sn'])
        
    # get english name
    nstart=fi+1
    fi=txt.find(') ', nstart)
    if(fi != -1):
        en = txt[nstart:fi].strip()
        lst = []
        enlist = en.split()
        for e in enlist:
            if(e.istitle()):
                lst.append(str.lower(e))
            else:
                lst.append(e)
        dtxt['en'] = ' '.join(lst)
        dtxt['eu'] = str.upper(dtxt['en'])
        
    # get chinese name
    nstart=fi+1
    while(str.isspace(txt[nstart])):
        nstart += 1
        
    fi=txt.find(' ', nstart)
    if(fi!=-1):
        dtxt['cn'] = txt[nstart:fi].strip()
    
    nstart = fi+1
    dtxt['txt'] = txt[nstart:].strip()
    
    df = df.append(dtxt, ignore_index=True)


# find the dumplicated
vcount = df.eu.value_counts()
print(vcount[vcount>1])

#doc.paragraphs.clear()
#print(df.index)
#df.sort_values(by='eu', inplace=True)

# columns=[sortname, title, description]
newdf = pd.DataFrame(columns=['sn','title','txt'])
dnew={}
for i in range(len(df)):
    # get full name
    rw = df.iloc[i]
    dnew['sn'] = rw.eu
    if(len(rw.sn)>0):
        dnew['title'] = "{en} ({sn})  {cn}  ".format(en=rw.en, sn=rw.sn, cn=rw.cn)
    else:
        dnew['title'] = "{en}  {cn}  ".format(en=rw.en, cn=rw.cn)
    dnew['txt'] = rw.txt
    newdf = newdf.append(dnew, ignore_index=True)
    
    # get short name
    su = rw.su.replace('.', '')
    if(len(su)>=3 and len(su)<6):
        dnew['sn'] = su
        if(str.isdigit(rw.sn[0])):
            dnew['sn'] = "ZZ" + su   # 排在字母后面
        dnew['title'] = "{sn}  {cn}  ".format(sn=rw.sn, cn=rw.cn)
        dnew['txt'] = "{en}的缩写。".format(en=rw.en)
        newdf = newdf.append(dnew, ignore_index=True)
        
print("Total: ", len(newdf), ", Full: ", len(df), ", Short: ", len(newdf)-len(df))

# save to docx
hd = ''
doc=docx.Document()
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = docx.shared.Pt(12)
sty._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
newdf.sort_values(by='sn', inplace=True)
for i in range(len(newdf)):
    rw=newdf.iloc[i]
    if (hd != rw.sn[0]):
        if(len(hd)>0):
            doc.add_page_break()
        hd = rw.sn[0]
        doc.add_heading(hd)
    p=doc.add_paragraph(style=sty)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = docx.shared.Pt(4)
    p.paragraph_format.space_after = 0
    run=p.add_run(rw.title)
    run.bold = True
    p.add_run(rw.txt)
    
    
doc.save(os.path.join(dir,"dict-new.docx"))

'''
doc=docx.Document()
for i in range(len(df)):
    rw = df.iloc[i]
    if(len(rw.sn)>0):
        txt = "{en}({sn})  {cn}  {txt}".format(en=rw.en, sn=rw.sn, cn=rw.cn, txt=rw.txt)
    else:
        txt = "{en}  {cn}  {txt}".format(en=rw.en, cn=rw.cn, txt=rw.txt)
    doc.add_paragraph(txt)    
    #print(rw.en)
    
doc.save("F:/Dict-book/dict-full.docx")

df.sort_values(by='su', inplace=True)
doc=docx.Document()
#doc.styles['Normal'].font.name = u'宋体'
for i in range(len(df)):
    rw = df.iloc[i]
    if(len(rw.sn)>=3 and len(rw.sn)<6):
        p = doc.add_paragraph()
        run = p.add_run("{sn}  {cn}  ".format(sn=rw.sn, cn=rw.cn))
        run.bold = True
        p.add_run("{en}的缩写".format(en=rw.en))
        
doc.save("F:/Dict-book/dict-short.docx")
'''