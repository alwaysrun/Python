# -*- coding: utf-8 -*-
"""

first：
pip install python_docx

读取原文档: abrasive flow machining (AFM)  流液研磨加工  是一种...。
拆分出‘英文名’、‘简写’、‘中文名’、‘注释’然后组合成所需样式：
‘英文名’(‘简写’)  ‘中文名’  ‘注释’
‘简写’  ‘中文名’  ‘英文名’的缩写
最后统一进行排序

@author: guodxu@qq.com
"""

import os
import re
import docx
import pandas as pd

dir=os.path.split(os.path.realpath(__file__))[0]
print(dir)
doc=docx.Document(os.path.join(dir,"dict-full.docx"))

print("paragraphs: ", len(doc.paragraphs))


docSave=docx.Document()
testCount=0

dtxt={}
# columns[shortname, shortnameUpper, englisname, englisnameUpper, chinesename, description]
df = pd.DataFrame(columns=['sn','su','en','eu','cn','txt'])
for par in doc.paragraphs:
    txt = par.text
    if(len(txt)<5):
        continue
    if(txt[-4:] == '的缩写。'):  # remove the short-entry
        continue
    
    title=''
    for rn in par.runs:
        if(rn.bold):
            title += rn.text
        else:
            break
    if(len(title)<2):
        print('No title:', txt)
        continue

    ## txt
    txt = txt[len(title):].strip()
    
    ## get cn & sn
    title = title.strip()
    rn = title.rfind('  ')
    cn=''
    if(rn==-1):
        print('No CN:', title)
    else:
        cn=title[rn:].strip()
        title=title[:rn].strip()
    
    sn=''
    rn = title.rfind('(')
    if(rn!=-1):
        sn=title[rn+1:].rstrip(')').strip()
        title=title[:rn].strip()
        if(re.search('[a-z]', sn)):
            print(rn, sn)
        
     
    if(len(title)<2):
        print('No En-title:', txt)

    dtxt['sn']=sn
    dtxt['su']=sn.upper()
    dtxt['en']=title
    dtxt['eu']=title.upper()
    dtxt['cn']=cn
    dtxt['txt']=txt
    df = df.append(dtxt, ignore_index=True)
        
#    testCount += 1
#    if(testCount>10): break    

#print('total: ', testCount)
print('total: ', len(df))

# find the dumplicated
vcount = df.eu.value_counts()
print(vcount[vcount>1])

#
doc.paragraphs.clear()
print(df.index)
df.sort_values(by='eu', inplace=True)

#
# columns=[sortname, title, description]
newdf = pd.DataFrame(columns=['sn','title','txt'])
dnew={}
for i in range(len(df)):
    # get full name
    rw = df.iloc[i]
    dnew['sn'] = rw.eu.replace('.', '').replace('-', '').replace('/', '')
    if(str.isdigit(dnew['sn'][0])):
        dnew['sn'] = 'zz' + dnew['sn']
    if(len(rw.sn)>0):
        dnew['title'] = "{en} ({sn})  {cn}  ".format(en=rw.en, sn=rw.sn, cn=rw.cn)
    else:
        dnew['title'] = "{en}  {cn}  ".format(en=rw.en, cn=rw.cn)
    dnew['txt'] = rw.txt
    newdf = newdf.append(dnew, ignore_index=True)
    
    # get short name
    su = rw.su.replace('.', '').replace('-', '').replace('/', '')
    if(len(su)>=3 and len(su)<6):
        dnew['sn'] = su
        if(not str.isalpha(rw.sn[0])):
            dnew['sn'] = "zz" + su   # 排在字母后面
        dnew['title'] = "{sn}  {cn}  ".format(sn=rw.sn, cn=rw.cn)
        dnew['txt'] = "{en}的缩写。".format(en=rw.en)
        newdf = newdf.append(dnew, ignore_index=True)
        
print("Total: ", len(newdf), ", Full: ", len(df), ", Short: ", len(newdf)-len(df))

# save to docx
hd = ''
doc=docx.Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
sty = doc.styles['Normal']
sty.font.size = docx.shared.Pt(12)
#sty.font.name = 'Times New Roman'
#sty._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
newdf.sort_values(by='sn', inplace=True)
for i in range(len(newdf)):
    rw=newdf.iloc[i]
    if (hd != rw.sn[0]):
        if(len(hd)>0):
            doc.add_page_break()
        hd = rw.sn[0]
        rtitle=None
        if(hd=='z'):
            rtitle = doc.add_heading('').add_run('以数字、希文字母起首的辞条')
        else:
            rtitle = doc.add_heading('').add_run(hd)
        rtitle.font.name = 'Times New Roman'
        rtitle._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
        
    p=doc.add_paragraph(style=sty)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = docx.shared.Pt(4)
    p.paragraph_format.space_after = 0
    run=p.add_run(rw.title)
    run.bold = True
    p.add_run(rw.txt)
    
    
doc.save(os.path.join(dir,"dictFull-new.docx"))

'''
# save short
doc=docx.Document()

df.sort_values(by='su', inplace=True)
doc=docx.Document()
#doc.styles['Normal'].font.name = u'宋体'
for i in range(len(df)):
    rw = df.iloc[i]
    if(len(rw.sn)>=3 and len(rw.sn)<6):
        p = doc.add_paragraph()
        run = p.add_run("{sn}  {cn}  ".format(sn=rw.sn, cn=rw.cn))
        run.bold = True
        p.add_run("{en}的缩写".format(en=rw.en))
        
doc.save(os.path.join(dir,"dict-short.docx"))
'''